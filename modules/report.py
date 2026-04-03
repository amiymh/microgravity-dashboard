"""Word methods report generator using python-docx."""

import io
import re
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _extract(log_lines: list[str], pattern: str, default: str = "") -> str:
    """Extract the first regex match from log lines."""
    for line in log_lines:
        m = re.search(pattern, line, re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return default


def get_reproduction_steps(
    tab_name: str,
    log_lines: list[str],
    gene_lists: dict[str, list[str]] | None = None,
) -> list[str]:
    """Generate numbered step-by-step reproduction instructions for a tab.

    Extracts actual parameter values from the analysis log and gene lists
    so a scientist can independently verify every result outside the tool.
    """
    gl = gene_lists or {}
    steps: list[str] = []

    if tab_name == "Volcano Plot":
        padj = _extract(log_lines, r"p-value threshold:\s*([\d.eE-]+)", "0.05")
        fc = _extract(log_lines, r"log2FoldChange\| threshold:\s*([\d.]+)", "1.0")
        n_up = _extract(log_lines, r"Upregulated.*?:\s*(\d+)", "?")
        n_down = _extract(log_lines, r"Downregulated.*?:\s*(\d+)", "?")
        steps = [
            "1. Open the filtered gene list CSV exported from the Volcano Plot tab, or use the gene list in Appendix A ('Volcano Plot').",
            f"2. Apply the classification logic to your DESeq2 results: a gene is Upregulated if log2FoldChange > {fc} AND padj < {padj}; Downregulated if log2FoldChange < -{fc} AND padj < {padj}; otherwise Not Significant.",
            "3. In Excel: add a helper column with =IF(AND(log2FC>{fc}, padj<{padj}), \"Up\", IF(AND(log2FC<-{fc}, padj<{padj}), \"Down\", \"NS\")). In R: use dplyr::case_when() with the same thresholds.".format(fc=fc, padj=padj),
            f"4. Count the results. You should get approximately {n_up} upregulated and {n_down} downregulated genes, matching the Analysis Details above.",
            "5. To recreate the plot in R: use EnhancedVolcano (https://bioconductor.org/packages/EnhancedVolcano/) or ggplot2 with geom_point(); in Python: use plotly or matplotlib scatter with x=log2FC, y=-log10(padj).",
        ]

    elif tab_name == "PCA Plot":
        var1 = _extract(log_lines, r"PC1.*?:\s*([\d.]+)%", "?")
        var2 = _extract(log_lines, r"PC2.*?:\s*([\d.]+)%", "?")
        n_genes = _extract(log_lines, r"genes used.*?:\s*(\d+)", "?")
        n_removed = _extract(log_lines, r"removed.*?:\s*(\d+)", "0")
        steps = [
            "1. Open your DESeq2 results file and extract the 16 sample columns: Earth_3D2, Earth_4D2, Earth_3E2, Earth_4E2, Earth_1D2, Earth_2D2, Earth_1E2, Earth_2E2, Space_3D2, Space_4D2, Space_3E2, Space_4E2, Space_1D2, Space_2D2, Space_1E2, Space_2E2.",
            "2. Transpose the matrix so that the 16 samples are rows and genes are columns.",
            f"3. Remove genes with zero variance across all 16 samples ({n_removed} genes were removed in this session, leaving {n_genes} genes).",
            "4. Apply StandardScaler normalization (mean=0, std=1 per gene). In Python: sklearn.preprocessing.StandardScaler() (https://scikit-learn.org/stable/modules/generated/sklearn.preprocessing.StandardScaler.html). In R: scale().",
            "5. Run PCA with 3 components. In Python: sklearn.decomposition.PCA(n_components=3) (https://scikit-learn.org/stable/modules/generated/sklearn.decomposition.PCA.html). In R: prcomp(scaled_data, center=FALSE, scale.=FALSE) since data is already scaled.",
            f"6. PC1 should explain approximately {var1}% of variance and PC2 approximately {var2}%. Small differences (<1%) are expected due to floating-point precision.",
        ]

    elif tab_name == "Top DEGs Heatmap":
        n_genes = _extract(log_lines, r"(\d+) genes selected", "50")
        steps = [
            f"1. From your filtered DEG list, select the top {n_genes} genes ranked by significance score (= |log2FoldChange| × -log10(padj)).",
            "2. Extract the 16 sample columns (8 Earth + 8 Space) for those genes.",
            "3. Z-score normalize each gene (row) across the 16 samples: z = (x - mean) / std. In Python: scipy.stats.zscore(row). In R: scale(t(matrix)).",
            "4. Optionally cluster genes using hierarchical clustering (ward linkage, Euclidean distance). In Python: scipy.cluster.hierarchy.linkage(z_matrix, method='ward'). In R: hclust(dist(z_matrix), method='ward.D2').",
            "5. Plot the heatmap using a diverging blue-white-red color scale centered at z=0. Each cell represents a gene's expression in a specific sample relative to its mean across all samples.",
        ]

    elif tab_name == "Pathway Enrichment":
        library = _extract(log_lines, r"Gene set library:\s*(.+)", "KEGG_2021_Human")
        n_genes = _extract(log_lines, r"Gene list submitted:\s*(\d+)", "?")
        direction = _extract(log_lines, r"Direction filter:\s*(.+)", "All")
        n_results = _extract(log_lines, r"Results returned:\s*(\d+)", "?")
        list_id = _extract(log_lines, r"Enrichr list ID:\s*(.+)", "")
        gene_key = next((k for k in gl if "pathway" in k.lower() or "enrichment" in k.lower()), None)
        gene_preview = ", ".join(gl[gene_key][:20]) + "..." if gene_key and gl[gene_key] else "(see Appendix A)"
        steps = [
            "1. Go to https://maayanlab.cloud/Enrichr/",
            f"2. In the text box, paste the gene list ({n_genes} genes, direction filter: {direction}). Gene list: {gene_preview}. Full list in Appendix A.",
            "3. Click 'Submit'.",
            f"4. On the results page, select the gene set library: {library}.",
            "5. Click on the library name to expand results. Results are ranked by combined score.",
            f"6. You should see approximately {n_results} enriched pathways. Compare the top pathway names, p-values, and overlap counts with the table in this report.",
            f"7. If an Enrichr list ID was returned ({list_id or 'N/A'}), you can revisit results at https://maayanlab.cloud/Enrichr/enrich?dataset={list_id}" if list_id else "7. Note: Enrichr list IDs are temporary and may expire.",
        ]

    elif tab_name == "Therapeutic Targets":
        n_queried = _extract(log_lines, r"Genes queried:\s*(\d+)", "200")
        n_found = _extract(log_lines, r"Genes with drug interactions found:\s*(\d+)", "?")
        n_drugs = _extract(log_lines, r"Unique drugs found:\s*(\d+)", "?")
        steps = [
            f"1. Go to https://dgidb.org/ and click 'Search Interactions'.",
            f"2. Paste the top {n_queried} gene symbols (by significance score) from Appendix A into the gene search box. Submit in batches of 20 if needed.",
            f"3. DGIdb should return drug interactions for approximately {n_found} genes, totaling approximately {n_drugs} unique drugs. Compare drug names and interaction types with the table in this report.",
            "4. For clinical stage data: go to https://platform.opentargets.org/",
            "5. Search for each gene that had drug interactions. Click on the gene, then go to the 'Known drugs' or 'Pharmacology' section.",
            "6. The 'Clinical Stage' column in this report shows values like PHASE_2, PHASE_3 — these come directly from the OpenTargets 'drugAndClinicalCandidates' field. Only values explicitly returned by the API are shown.",
        ]

    elif tab_name == "Biomarker Discovery":
        n_filtered = _extract(log_lines, r"Genes passing.*filter:\s*(\d+)", "?")
        n_processed = _extract(log_lines, r"Total genes processed.*:\s*(\d+)", "?")
        n_with_assoc = _extract(log_lines, r"Genes with disease associations:\s*(\d+)", "?")
        example_line = ""
        for line in log_lines:
            if line.startswith("Example:"):
                example_line = line
                break
        steps = [
            "1. Go to https://platform.opentargets.org/",
            f"2. For each of the {n_processed} genes (from Appendix A), search the gene symbol and go to the 'Associated Diseases' tab.",
            "3. Note the top 5 disease associations and their scores (0-1 scale).",
            f"4. Approximately {n_with_assoc} genes should have at least one disease association with score > 0.",
            "5. Compute the Biomarker Score for each gene: |log2FoldChange| × -log10(padj) × (disease_score + 0.1).",
            f"6. {example_line}" if example_line else "6. Compare your computed scores with the Biomarker Score column in the report table.",
            "7. Rank genes by Biomarker Score descending — the order should match the report table.",
        ]

    elif tab_name == "Cytotoxicity & Apoptosis":
        pathway_lines = [l for l in log_lines if l.strip().startswith(("Apoptosis", "TNF", "p53", "Reactive", "Inflammatory", "Gene set"))]
        contingency_lines = [l for l in log_lines if "overlap=" in l]
        steps = [
            "1. Go to https://www.gsea-msigdb.org/gsea/msigdb/human/search.jsp",
            "2. Search for each hallmark gene set used in this analysis:",
        ]
        for pl in pathway_lines:
            steps.append(f"   - {pl}")
        steps.append("3. Download each gene set (click the set name, then 'Gene Symbols' or JSON download). Verify the gene count matches the 'Pathway Size' in the report table.")
        steps.append("4. Intersect each gene set with your DEG list (the set of genes that passed your padj and log2FC filters).")
        if contingency_lines:
            steps.append("5. For each pathway, verify the Fisher's exact test using these values:")
            for cl in contingency_lines:
                steps.append(f"   - {cl.strip()}")
            steps.append("   Fisher's exact test contingency table: [[overlap, pathway_only], [deg_only, background]]. Background = 20,000 - overlap - pathway_only - deg_only. Use scipy.stats.fisher_exact() in Python or fisher.test() in R with alternative='greater'.")
        else:
            steps.append("5. Run Fisher's exact test for each pathway. Contingency table: [[overlap, pathway_only], [deg_only, 20000-overlap-pathway_only-deg_only]], alternative='greater'.")
        steps.append("6. Compare overlap counts and p-values with the report table.")
        # Re-number
        steps = [f"{i+1}. {s.split('. ', 1)[1]}" if s[0].isdigit() and '. ' in s else s for i, s in enumerate(steps)]

    elif tab_name == "Disease Cross-Reference":
        disease = _extract(log_lines, r"Disease searched:\s*'?([^']+)'?", "?")
        efo_id = _extract(log_lines, r"EFO ID:\s*(\S+)", "?")
        threshold = _extract(log_lines, r"score threshold:\s*([\d.]+)", "0.3")
        n_overlap = _extract(log_lines, r"Overlapping genes found:\s*(\d+)", "?")
        n_disease = _extract(log_lines, r"Disease genes returned.*:\s*(\d+)", "?")
        overlap_key = next((k for k in gl if "disease" in k.lower() or "overlap" in k.lower()), None)
        overlap_genes = ", ".join(gl[overlap_key][:15]) if overlap_key and gl[overlap_key] else "(see Appendix A)"
        steps = [
            "1. Go to https://platform.opentargets.org/",
            f"2. Search for disease: '{disease}'. Select the result matching EFO ID: {efo_id}.",
            "3. Go to the 'Associated Targets' tab on the disease page.",
            f"4. Set the association score filter to >= {threshold}. Approximately {n_disease} genes should appear.",
            "5. Export or note the gene list from OpenTargets.",
            f"6. Intersect this gene list with your DEG list. You should find approximately {n_overlap} overlapping genes.",
            f"7. The overlapping genes should include: {overlap_genes}.",
            "8. Compare association scores for each overlapping gene with the 'Disease Association Score' column in the report table.",
        ]

    else:
        steps = [
            "1. Refer to the Analysis Details section above for exact parameters.",
            "2. See Appendix A for gene lists used in this analysis.",
        ]

    return steps


def generate_report(
    session_id: str = "",
    filename: str = "unknown",
    sheet_name: str = "SDEGs",
    gene_count: int = 0,
    analysis_logs: dict[str, list[str]] | None = None,
    gene_lists: dict[str, list[str]] | None = None,
    versions: dict | None = None,
) -> bytes:
    """Generate a complete Word methods report.

    Args:
        session_id: UUID of the current session.
        filename: Name of the uploaded file.
        sheet_name: Which Excel sheet was used.
        gene_count: Total genes loaded.
        analysis_logs: Dict of {tab_name: [log_lines]} for each analysis run.
        gene_lists: Dict of {analysis_name: [gene_symbols]} for Appendix A.
        versions: Dict of library/API versions for Appendix C.

    Returns:
        bytes of the .docx file.
    """
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading("Analysis Methods Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Microgravity RNA Dashboard", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    info_items = [
        ("Date Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("File Analyzed", filename),
        ("Sheet", sheet_name),
        ("Total Genes Loaded", str(gene_count)),
        ("Session ID", session_id or "N/A"),
    ]
    table = doc.add_table(rows=len(info_items), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(info_items):
        table.cell(i, 0).text = key
        table.cell(i, 1).text = val

    doc.add_page_break()

    # --- Analysis Sections ---
    all_tabs = [
        "Volcano Plot", "PCA Plot", "Top DEGs Heatmap",
        "Pathway Enrichment", "Therapeutic Targets",
        "Biomarker Discovery", "Cytotoxicity & Apoptosis",
        "Disease Cross-Reference",
    ]

    logs = analysis_logs or {}
    gl = gene_lists or {}

    for tab_name in all_tabs:
        doc.add_heading(tab_name, level=1)

        if tab_name in logs and logs[tab_name]:
            doc.add_heading("Analysis Details", level=2)
            for line in logs[tab_name]:
                doc.add_paragraph(line, style="List Bullet")

            doc.add_heading("To Reproduce Manually", level=2)
            repro_steps = get_reproduction_steps(tab_name, logs[tab_name], gl)
            for step in repro_steps:
                doc.add_paragraph(step)
        else:
            doc.add_paragraph("Not run in this session.", style="Intense Quote")

    doc.add_page_break()

    # --- Appendix A: Gene Lists ---
    doc.add_heading("Appendix A: Gene Lists", level=1)
    if gl:
        for name, genes in gl.items():
            doc.add_heading(name, level=2)
            doc.add_paragraph(", ".join(genes[:200]))
            if len(genes) > 200:
                doc.add_paragraph(f"... and {len(genes) - 200} more genes")
    else:
        doc.add_paragraph("No gene lists recorded in this session.")

    # --- Appendix B: API Responses ---
    doc.add_heading("Appendix B: API Response Summaries", level=1)
    doc.add_paragraph("API response details are included in each analysis section's log above.")

    # --- Appendix C: Versions ---
    doc.add_heading("Appendix C: Software and Database Versions", level=1)
    versions = versions or {}
    if versions:
        table = doc.add_table(rows=len(versions), cols=2)
        table.style = "Table Grid"
        for i, (key, val) in enumerate(versions.items()):
            table.cell(i, 0).text = str(key)
            table.cell(i, 1).text = str(val)
    else:
        doc.add_paragraph("Version information not available.")

    # --- Appendix D: References ---
    doc.add_heading("Appendix D: Method References", level=1)
    refs = [
        "DESeq2: Love MI et al. Genome Biology 2014. DOI: 10.1186/s13059-014-0550-8",
        "Enrichr: Kuleshov MV et al. Nucleic Acids Res 2016. DOI: 10.1093/nar/gkw377",
        "OpenTargets: Ochoa D et al. Nucleic Acids Res 2023. DOI: 10.1093/nar/gkac1033",
        "DGIdb: Freshour SL et al. Nucleic Acids Res 2021. DOI: 10.1093/nar/gkaa1084",
        "MSigDB: Liberzon A et al. Cell Systems 2015. DOI: 10.1016/j.cels.2015.12.004",
        "Fisher's Exact Test: Fisher RA. J Royal Stat Soc 1922.",
        "PCA: Jolliffe IT. Principal Component Analysis. Springer 2002.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
