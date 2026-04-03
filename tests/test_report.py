"""Tests for Word methods report module."""

import os
import sys
import re
import io
import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from modules.report import generate_report, get_reproduction_steps


# Realistic analysis logs mimicking what the dashboard actually produces
SAMPLE_LOGS = {
    "Volcano Plot": [
        "Total genes in dataset: 4522",
        "Adjusted p-value threshold: 0.05",
        "|log2FoldChange| threshold: 1.0",
        "Genes passing padj threshold: 4522",
        "Genes passing |log2FC| threshold: 2100",
        "Upregulated (both thresholds): 1100",
        "Downregulated (both thresholds): 1000",
        "Not significant: 2422",
    ],
    "PCA Plot": [
        "Number of samples used: 16",
        "Number of genes used: 4400",
        "Genes removed (zero variance): 122",
        "PC1 variance explained: 32.5%",
        "PC2 variance explained: 18.1%",
        "PC3 variance explained: 9.2%",
    ],
    "Pathway Enrichment": [
        "Gene list submitted: 1100 genes",
        "Direction filter: Upregulated",
        "Gene set library: KEGG_2021_Human",
        "API endpoint: https://maayanlab.cloud/Enrichr",
        "Enrichr list ID: 12345",
        "Timestamp: 2026-04-03T12:00:00",
        "Data source: Enrichr API (live)",
        "Results returned: 20 pathways",
    ],
    "Therapeutic Targets": [
        "Genes queried: 200 (top by significance score)",
        "DGIdb API endpoint: https://dgidb.org/api/v2/interactions.json",
        "DGIdb batch size: 20 genes per request",
        "OpenTargets API endpoint: https://api.platform.opentargets.org/api/v4/graphql",
        "Timestamp: 2026-04-03T12:01:00",
        "Genes with drug interactions found: 45",
        "Unique drugs found: 120",
    ],
    "Biomarker Discovery": [
        "Genes passing |log2FC| >= 2.0 filter: 800",
        "Total genes processed (no cap): 800",
        "Batch size: 20 genes per batch",
        "API endpoint: https://api.platform.opentargets.org/api/v4/graphql",
        "Timestamp: 2026-04-03T12:02:00",
        "Genes with disease associations: 350",
        "Genes with no associations found: 450",
        "Example: ORM1 score = |4.25| × -log10(1.76e-121) × (0.85 + 0.1)",
    ],
    "Cytotoxicity & Apoptosis": [
        "MSigDB API endpoint: https://www.gsea-msigdb.org/gsea/msigdb/human/download_geneset.jsp",
        "Timestamp: 2026-04-03T12:03:00",
        "Gene set 'Apoptosis': 161 genes (MSigDB API)",
        "Gene set 'TNF Signaling via NF-kB': 200 genes (MSigDB API)",
        "  Apoptosis: overlap=42, pathway_only=119, Fisher p=1.23e-08",
        "  TNF Signaling via NF-kB: overlap=55, pathway_only=145, Fisher p=3.45e-12",
    ],
    "Disease Cross-Reference": [
        "Disease searched: 'Alzheimer disease'",
        "Disease EFO ID: EFO_0000249",
        "Association score threshold: 0.3",
        "Direction filter: All",
        "API endpoint: https://api.platform.opentargets.org/api/v4/graphql",
        "Timestamp: 2026-04-03T12:04:00",
        "Disease genes returned by OpenTargets: 200",
        "DEGs queried: 2100",
        "Overlapping genes found: 15",
    ],
}

SAMPLE_GENE_LISTS = {
    "Volcano Plot": ["ORM1", "TNF", "IL6", "TP53", "BCL2"],
    "Pathway Enrichment": ["TNF", "IL6", "CASP3", "BCL2", "NFKB1"],
    "Disease Cross-Reference overlap": ["TNF", "IL6", "TP53"],
}


class TestGetReproductionSteps:
    """Test the step generator directly."""

    @pytest.mark.parametrize("tab_name", list(SAMPLE_LOGS.keys()))
    def test_returns_list_of_strings(self, tab_name):
        steps = get_reproduction_steps(tab_name, SAMPLE_LOGS[tab_name], SAMPLE_GENE_LISTS)
        assert isinstance(steps, list)
        assert all(isinstance(s, str) for s in steps)

    @pytest.mark.parametrize("tab_name", list(SAMPLE_LOGS.keys()))
    def test_has_numbered_steps(self, tab_name):
        steps = get_reproduction_steps(tab_name, SAMPLE_LOGS[tab_name], SAMPLE_GENE_LISTS)
        numbered = [s for s in steps if re.match(r"^\d+\.", s)]
        assert len(numbered) >= 2, f"{tab_name}: expected numbered steps, got {steps}"

    @pytest.mark.parametrize("tab_name", list(SAMPLE_LOGS.keys()))
    def test_contains_url(self, tab_name):
        steps = get_reproduction_steps(tab_name, SAMPLE_LOGS[tab_name], SAMPLE_GENE_LISTS)
        all_text = " ".join(steps)
        assert "https://" in all_text, f"{tab_name}: no URL found in reproduction steps"

    def test_volcano_has_threshold_values(self):
        steps = get_reproduction_steps("Volcano Plot", SAMPLE_LOGS["Volcano Plot"])
        text = " ".join(steps)
        assert "0.05" in text
        assert "1.0" in text
        assert "1100" in text  # upregulated count

    def test_pathway_has_library_name(self):
        steps = get_reproduction_steps("Pathway Enrichment", SAMPLE_LOGS["Pathway Enrichment"], SAMPLE_GENE_LISTS)
        text = " ".join(steps)
        assert "KEGG_2021_Human" in text
        assert "maayanlab.cloud" in text

    def test_disease_has_efo_id(self):
        steps = get_reproduction_steps("Disease Cross-Reference", SAMPLE_LOGS["Disease Cross-Reference"], SAMPLE_GENE_LISTS)
        text = " ".join(steps)
        assert "EFO_0000249" in text
        assert "Alzheimer" in text

    def test_biomarker_has_formula(self):
        steps = get_reproduction_steps("Biomarker Discovery", SAMPLE_LOGS["Biomarker Discovery"])
        text = " ".join(steps)
        assert "log2FoldChange" in text or "log2FC" in text
        assert "disease_score" in text

    def test_cytotoxicity_has_fisher_details(self):
        steps = get_reproduction_steps("Cytotoxicity & Apoptosis", SAMPLE_LOGS["Cytotoxicity & Apoptosis"])
        text = " ".join(steps)
        assert "fisher" in text.lower() or "Fisher" in text
        assert "msigdb" in text.lower()

    def test_pca_has_variance(self):
        steps = get_reproduction_steps("PCA Plot", SAMPLE_LOGS["PCA Plot"])
        text = " ".join(steps)
        assert "32.5" in text  # PC1 variance
        assert "StandardScaler" in text or "scale" in text.lower()

    def test_targets_has_dgidb_url(self):
        steps = get_reproduction_steps("Therapeutic Targets", SAMPLE_LOGS["Therapeutic Targets"])
        text = " ".join(steps)
        assert "dgidb.org" in text
        assert "opentargets.org" in text


class TestGenerateReport:
    def test_generates_without_error(self):
        data = generate_report(session_id="test-123", filename="test.xlsx", gene_count=100)
        assert isinstance(data, bytes)
        assert len(data) > 0

    def test_is_valid_docx(self):
        data = generate_report()
        from docx import Document
        doc = Document(io.BytesIO(data))
        assert len(doc.paragraphs) > 0

    def test_contains_expected_headings(self):
        data = generate_report(analysis_logs=SAMPLE_LOGS)
        from docx import Document
        doc = Document(io.BytesIO(data))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        heading_text = " ".join(headings)
        assert "Volcano Plot" in heading_text
        assert "Appendix A" in heading_text
        assert "Appendix C" in heading_text

    def test_to_reproduce_heading_for_each_run_tab(self):
        data = generate_report(analysis_logs=SAMPLE_LOGS)
        from docx import Document
        doc = Document(io.BytesIO(data))
        headings = [p.text for p in doc.paragraphs if "Heading" in (p.style.name or "")]
        reproduce_count = sum(1 for h in headings if "To Reproduce Manually" in h)
        assert reproduce_count == len(SAMPLE_LOGS), (
            f"Expected {len(SAMPLE_LOGS)} 'To Reproduce Manually' headings, got {reproduce_count}"
        )

    def test_numbered_steps_in_report(self):
        data = generate_report(analysis_logs=SAMPLE_LOGS, gene_lists=SAMPLE_GENE_LISTS)
        from docx import Document
        doc = Document(io.BytesIO(data))
        full_text = " ".join(p.text for p in doc.paragraphs)
        # Should contain numbered steps like "1. Go to"
        assert re.search(r"\d+\.\s+(Go to|Open|From|Apply|Extract|Search|In |Paste|For |Set )", full_text)

    def test_urls_in_report(self):
        data = generate_report(analysis_logs=SAMPLE_LOGS, gene_lists=SAMPLE_GENE_LISTS)
        from docx import Document
        doc = Document(io.BytesIO(data))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "https://maayanlab.cloud" in full_text
        assert "https://dgidb.org" in full_text
        assert "https://platform.opentargets.org" in full_text
        assert "https://www.gsea-msigdb.org" in full_text

    def test_gene_lists_in_appendix(self):
        data = generate_report(gene_lists=SAMPLE_GENE_LISTS)
        from docx import Document
        doc = Document(io.BytesIO(data))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "TP53" in full_text

    def test_versions_in_appendix(self):
        data = generate_report(versions={"Python": "3.14.0", "pandas": "2.2.0"})
        from docx import Document
        doc = Document(io.BytesIO(data))
        found = False
        for table in doc.tables:
            for row in table.rows:
                if "Python" in [c.text for c in row.cells]:
                    found = True
        assert found

    def test_not_run_tabs_noted(self):
        data = generate_report()
        from docx import Document
        doc = Document(io.BytesIO(data))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Not run" in full_text
